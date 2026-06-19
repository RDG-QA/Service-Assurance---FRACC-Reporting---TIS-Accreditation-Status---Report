#!/usr/bin/env python3
"""
====================================================================
FRACC TIS Accreditation Status Report Generator
====================================================================
Rail Delivery Group — Service Assurance
Generates the monthly FRACC paper and supporting FRACC Excel data
file from a collated TOC TIS weekly report.

Usage:
    python generate_fracc_paper.py \\
        --collated  "TOC TIS Accreditation Status - Collated Report - v3.0.xlsx" \\
        --template  "FRACC Paper Template.docx" \\
        --paper-date "16 June 2026" \\
        [--output-dir "./output"]
        [--news-image "assist_news.png"]

Supported collated file formats
    1. Raw stacked format  — multiple weekly blocks in one sheet,
       each starting with "Shift date range used: From …" row,
       with a "Lennon profit centre" column header (first block only).
       This is the output of the portal downloader collate.py script.

    2. Pivot export format — a sheet that already has an
       "Owning Group" column (e.g. the Pivot Data sheet of a
       previously generated FRACC Excel output).

Parameters:
    --collated      Path to the collated master Excel report.
    --template      Path to the FRACC Word template (.docx).
    --paper-date    Date string for the Paper Date field, e.g. "16 June 2026".
    --output-dir    Optional output folder (default: current directory).
    --news-image    Optional path to ASSIST News widget screenshot (.png).

Outputs (written to --output-dir):
    FRACC - TIS Accreditation Status - Latest - <YYYYMMDD_HHMM>.xlsx
    FRACC Paper - Accreditation Status Update - <DD Month YYYY>.docx

Dependencies:
    pip install pandas openpyxl python-docx matplotlib lxml Pillow
====================================================================
"""

import argparse
import copy
import io
import os
import re
import sys
import zipfile
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


# ════════════════════════════════════════════════════════════════════
# CONSTANTS — Appendix / Owning Group structure
# ════════════════════════════════════════════════════════════════════

# Display order of Owning Groups in the paper (Appendix letter → display name)
APPENDIX_MAP = [
    ("A",  "Go Ahead"),
    ("B",  "Transport UK Group (formerly Abellio Group)"),
    ("C",  "Serco / Transport UK Group (formerly Abellio Group)"),
    ("D",  "First Group"),
    ("E",  "Directly Operated Railway"),
    ("F",  "Arriva"),
    ("G",  "Transport for Wales"),
    ("H",  "London Overground"),
    ("I",  "GTS Elizabeth Line"),
    ("K",  "Heathrow Express Operating Company"),
    ("L",  "Scottish Rail Holdings"),
]

# ── LPC → Owning Group ──────────────────────────────────────────────
# Definitive mapping of every Lennon Profit Centre (LPC) code that
# appears in the portal data to its Owning Group (Appendix group).
# Only LPCs listed here are included in the FRACC output.
# All others (3rd-party retailers, HQ inputs, etc.) are excluded.
LPC_TO_OWNING_GROUP = {
    # Go Ahead (currently no LPCs after GTR remap, kept for future use)
    # — (none) —

    # Transport UK Group (formerly Abellio Group)
    "EAST MIDLANDS RAILWAY":            "Transport UK Group (formerly Abellio Group)",

    # Serco / Transport UK Group (formerly Abellio Group)
    "MERSEYRAIL ELECTRICS 2002 LTD":    "Serco / Transport UK Group (formerly Abellio Group)",

    # First Group
    "GREAT WESTERN RAILWAY HK":         "First Group",
    "HULL TRAINS LTD":                  "First Group",
    "LUMO":                             "First Group",
    "LUMO STIRLING":                    "First Group",
    "WEST COAST PARTNERSHIP":           "First Group",

    # Directly Operated Railway
    "C2C":                              "Directly Operated Railway",
    "GREATER ANGLIA":                   "Directly Operated Railway",
    "GTR - SOUTHERN & GATWICK EXPRESS": "Directly Operated Railway",
    "GTR-THAMESLINK & GREAT NORTHERN":  "Directly Operated Railway",
    "ISLAND LINE":                      "Directly Operated Railway",
    "LONDON NORTH EASTERN RAILWAY":     "Directly Operated Railway",
    "NORTHERN - WEST":                  "Directly Operated Railway",
    "SOUTH WESTERN RAILWAY":            "Directly Operated Railway",
    "SOUTHEASTERN":                     "Directly Operated Railway",
    "TRANSPENNINE TRAINS":              "Directly Operated Railway",
    "WEST MIDLANDS TRAINS LTD":         "Directly Operated Railway",
    "WM TRAINS":                        "Directly Operated Railway",   # portal alias

    # Arriva
    "ARRIVA CROSS COUNTRY":             "Arriva",
    "GRAND CENTRAL":                    "Arriva",
    "THE CHILTERN RAILWAY CO. LTD":     "Arriva",

    # Transport for Wales
    "TRANSPORT FOR WALES":              "Transport for Wales",

    # London Overground
    "LONDON OVERGROUND":                "London Overground",

    # GTS Elizabeth Line
    "GTS ELIZABETH LINE":               "GTS Elizabeth Line",

    # Heathrow Express Operating Company
    "HEATHROW EXPRESS LTD":             "Heathrow Express Operating Company",

    # Scottish Rail Holdings
    "CALEDONIAN SLEEPERS":              "Scottish Rail Holdings",
    "SCOTRAIL":                         "Scottish Rail Holdings",
}

# Normalisation aliases for the Owning Group field when reading a pivot export
# (handles minor naming variations in previously generated files)
OG_ALIASES = {
    "serco/transport uk group (formerly abellio group)":
        "Serco / Transport UK Group (formerly Abellio Group)",
    "transport uk group (formerly abellio group)":
        "Transport UK Group (formerly Abellio Group)",
    "gts elizabeth line":   "GTS Elizabeth Line",
    "arriva":               "Arriva",
    "go ahead":             "Go Ahead",
    "first group":          "First Group",
    "directly operated railway": "Directly Operated Railway",
    "transport for wales":  "Transport for Wales",
    "london overground":    "London Overground",
    "heathrow express operating company": "Heathrow Express Operating Company",
    "scottish rail holdings": "Scottish Rail Holdings",
}

# Bookmark names in the Word template for each appendix
BOOKMARK_MAP = {
    "A": "Appendix_A",
    "B": "Appendix_B",
    "C": "Appendix_C",
    "D": "Appendix_D",
    "E": "Appendix_E",
    "F": "Appendix_F",
    "G": "Appendix_G",
    "H": "Appendix_H",
    "I": "Appendix_I",
    "K": "Appendix_K",
    "L": "Appendix_L",
    "M": "AppendixM",
    "N": "AppendixN",
    "O": "AppendixO",
    "P": "AppendixP",
}

# Accreditation state colour coding (Excel cell fill / font)
STATE_FILL = {
    "Accredited":               "C6EFCE",
    "Pilot phase":              "E2EFDA",
    "Accreditation expired":    "FFC7CE",
    "Application acknowledged": "FFEB9C",
}
STATE_FONT = {
    "Accredited":               "276221",
    "Pilot phase":              "375623",
    "Accreditation expired":    "9C0006",
    "Application acknowledged": "9C6500",
}

# Excel structural colours
HDR_FILL_HEX  = "1F3864"   # RDG navy — Owning Group header rows
HDR_FONT_HEX  = "FFFFFF"   # White
SUBHDR_FILL   = "D6E4F0"   # Light blue — TOC sub-header rows
SUBHDR_FONT   = "1F3864"   # RDG navy
COL_HDR_FILL  = "EBF3FB"   # Very light blue — column header rows
COL_HDR_FONT  = "1F3864"   # RDG navy


# ════════════════════════════════════════════════════════════════════
# STEP 1 — Load & normalise the collated report
# ════════════════════════════════════════════════════════════════════

def load_collated_report(path: str) -> pd.DataFrame:
    """
    Load the collated TOC TIS master report.

    Detects two formats automatically:
      A) Raw stacked format  — multiple "Shift date range…" blocks in
         one sheet.  Produced by collate.py / the weekly portal download.
      B) Pivot export format — sheet already has 'Owning Group' and
         'Count' columns (e.g. Pivot Data sheet of a previous FRACC Excel).

    Returns a normalised DataFrame with columns:
        Owning Group, TOC / LPC, System, Machine Type ID,
        Version, State, Count, CoA Expiry
    Only rows whose TOC/LPC appears in LPC_TO_OWNING_GROUP are kept.
    """
    print(f"\n[1] Loading collated report: {os.path.basename(path)}")

    # ── Try pivot export (scan all sheets for 'Owning Group' + 'Count') ──
    try:
        xl = pd.ExcelFile(path, engine="openpyxl")
        for sheet in xl.sheet_names:
            df = xl.parse(sheet)
            if "Owning Group" in df.columns and "Count" in df.columns:
                df = _normalise_pivot(df)
                print(f"    Detected pivot export (sheet='{sheet}') — "
                      f"{len(df)} rows before TOC filter")
                df = _apply_lpc_filter(df)
                print(f"    After TOC filter: {len(df)} rows")
                return df
    except Exception as e:
        print(f"    (pivot scan failed: {e})")

    # ── Fall back to raw stacked format ───────────────────────────────
    raw = pd.read_excel(path, header=None, engine="openpyxl")
    df  = _parse_raw_stacked(raw)
    print(f"    Detected raw stacked format — {len(df)} rows before TOC filter")
    df  = _apply_lpc_filter(df)
    print(f"    After TOC filter: {len(df)} rows")
    return df


def _parse_raw_stacked(raw: pd.DataFrame) -> pd.DataFrame:
    """
    Parse a raw stacked collated file (multiple weekly blocks in one sheet).

    Structure per block:
        "Shift date range used: From DD-Mon-YYYY to DD-Mon-YYYY"  ← block header
        [optional "Lennon profit centre" column header row]
        data rows …

    The column header row appears only in the first block; subsequent blocks
    jump straight into data rows.
    """
    # Find the column header row (Lennon profit centre)
    col_header_idx = None
    for i, row in raw.iterrows():
        if str(row.iloc[0]).strip() == "Lennon profit centre":
            col_header_idx = i
            break

    if col_header_idx is None:
        raise ValueError("Could not find 'Lennon profit centre' column header row "
                         "in the collated file.")

    col_names = [str(c).replace("\n", " ").strip()
                 for c in raw.iloc[col_header_idx].tolist()]

    # Find all block-start rows ("Shift date range…")
    block_starts = [
        i for i, row in raw.iterrows()
        if "Shift date range" in str(row.iloc[0])
    ]

    if not block_starts:
        raise ValueError("No 'Shift date range' rows found — "
                         "is this a raw stacked collated file?")

    frames = []
    for bi, start in enumerate(block_starts):
        end       = block_starts[bi + 1] if bi + 1 < len(block_starts) else len(raw)
        date_str  = str(raw.iloc[start, 0])

        block = raw.iloc[start + 1:end].copy()
        # Drop column-header rows that got re-embedded (first-block artefact)
        block = block[block.iloc[:, 0] != "Lennon profit centre"]
        block = block[block.iloc[:, 0].notna()]
        block = block[block.iloc[:, 0].astype(str).str.strip() != ""]
        block.columns = col_names
        block = block.reset_index(drop=True)

        m = re.search(r"to\s+(\d{2}-\w+-\d{4})", date_str)
        block["_report_date"] = (
            pd.to_datetime(m.group(1), format="%d-%b-%Y") if m else pd.NaT
        )
        frames.append(block)

    combined = pd.concat(frames, ignore_index=True)
    combined["Count"] = (
        pd.to_numeric(combined["Count"], errors="coerce").fillna(0).astype(int)
    )

    # Rename raw columns to standard names
    combined = combined.rename(columns={
        "Lennon profit centre": "TOC / LPC",
        "CoA expiry date":      "CoA Expiry",
    })

    # Derive Owning Group from LPC
    combined["Owning Group"] = combined["TOC / LPC"].apply(
        lambda lpc: LPC_TO_OWNING_GROUP.get(str(lpc).strip().upper(), None)
    )

    return combined[[
        "Owning Group", "TOC / LPC", "System", "Machine Type ID",
        "Version", "State", "Count", "CoA Expiry",
    ]]


def _normalise_pivot(df: pd.DataFrame) -> pd.DataFrame:
    """Normalise a pivot-style export that already has an Owning Group column."""
    df = df.copy()
    df["Count"] = pd.to_numeric(df["Count"], errors="coerce").fillna(0).astype(int)
    # Normalise Owning Group names
    df["Owning Group"] = df["Owning Group"].apply(_canonicalise_og)
    # Re-derive Owning Group from LPC for overridden TOCs
    df["Owning Group"] = df.apply(
        lambda r: LPC_TO_OWNING_GROUP.get(
            str(r.get("TOC / LPC", "")).strip().upper(), r["Owning Group"]
        ),
        axis=1,
    )
    return df[[c for c in [
        "Owning Group", "TOC / LPC", "System", "Machine Type ID",
        "Version", "State", "Count", "CoA Expiry",
    ] if c in df.columns]]


def _canonicalise_og(name) -> str:
    if pd.isna(name):
        return str(name)
    key = str(name).strip().lower()
    return OG_ALIASES.get(key, str(name).strip())


def _apply_lpc_filter(df: pd.DataFrame) -> pd.DataFrame:
    """
    Keep only rows whose Owning Group is a known appendix group
    (i.e. rows with a valid LPC mapping).
    """
    known_ogs = {og for _, og in APPENDIX_MAP}
    return df[df["Owning Group"].isin(known_ogs)].reset_index(drop=True)


# ════════════════════════════════════════════════════════════════════
# STEP 2 — Apply peak-count logic
# ════════════════════════════════════════════════════════════════════

def apply_peak_counts(df: pd.DataFrame) -> pd.DataFrame:
    """
    For every (Owning Group, TOC/LPC, System, Machine Type ID, Version)
    combination, take the highest Count ever recorded across all historical
    weekly blocks.  State and CoA Expiry come from the peak-count row.
    """
    print("\n[2] Applying peak-count logic…")
    key_cols = ["Owning Group", "TOC / LPC", "System", "Machine Type ID", "Version"]
    peak = (
        df.sort_values("Count", ascending=False)
          .drop_duplicates(subset=key_cols, keep="first")
          .sort_values(["Owning Group", "TOC / LPC", "System",
                        "Machine Type ID", "Version"])
          .reset_index(drop=True)
    )
    print(f"    {len(df)} rows → {len(peak)} unique peak-count combinations")
    return peak


# ════════════════════════════════════════════════════════════════════
# EXCEL STYLE HELPERS
# ════════════════════════════════════════════════════════════════════

def _thin_border():
    s = Side(style="thin", color="BFBFBF")
    return Border(left=s, right=s, top=s, bottom=s)


def _hdr_font(bold=True, colour=HDR_FONT_HEX, size=11):
    return Font(bold=bold, color=colour, name="Calibri", size=size)


def _hdr_fill(hex_colour=HDR_FILL_HEX):
    return PatternFill("solid", fgColor=hex_colour)


def _centre():
    return Alignment(horizontal="center", vertical="center", wrap_text=True)


def _left():
    return Alignment(horizontal="left", vertical="center", wrap_text=True)


def _set_cell(cell, value=None, font=None, fill=None, alignment=None, border=None):
    if value is not None:
        cell.value = value
    if font:
        cell.font = font
    if fill:
        cell.fill = fill
    if alignment:
        cell.alignment = alignment
    if border:
        cell.border = border


# ════════════════════════════════════════════════════════════════════
# STEP 3 — Build FRACC Excel output
# ════════════════════════════════════════════════════════════════════

def build_fracc_excel(
    peak_df: pd.DataFrame,
    paper_date_str: str,
    output_dir: str,
    timestamp: str,
) -> tuple:
    """
    Build the FRACC TIS Accreditation Status Excel workbook.

    Formatting matches the reference file exactly:
      Font:        Aptos, 10pt (OG headers: 11pt)
      OG header:   fill=1F4E79, font=white, bold, center, row-height=18
                   border: left/top/bottom=medium, right=thin
      TOC header:  fill=2E75B6, font=white, bold, center, row-height=15
                   border: left=medium, right=thin, top/bottom=thin
                   merged A:F
      Col header:  fill=B4C6E7, no font colour override, bold, center, row-height=15
                   border: left=medium on A, medium on F, thin elsewhere
      Data rows:   entire row filled with state colour, center-aligned, row-height=14
                   border: left=medium on A, medium on F, thin elsewhere
      State fills: Accredited=C6EFCE, Pilot phase=FFEB9C, Expired=FFC7CE
      Spacer:      1 blank row between OG groups

    Sheet 1 — TIS Accreditation Status  (grouped view)
    Sheet 2 — Pivot Data                (flat table)
    Sheet 3 — Accreditation Status Chart (summary + source for pie)
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    print("\n[3] Building FRACC Excel file…")
    wb = Workbook()

    # ── Colour/style constants (from reference file) ─────────────────
    OG_FILL   = "1F4E79"   # dark navy
    OG_FONT   = "FFFFFF"
    TOC_FILL  = "2E75B6"   # mid blue
    TOC_FONT  = "FFFFFF"
    COL_FILL  = "B4C6E7"   # light blue

    STATE_FILL = {
        "Accredited":               "C6EFCE",
        "Pilot phase":              "FFEB9C",
        "Accreditation expired":    "FFC7CE",
        "Application acknowledged": "FFEB9C",
        "Abandoned":                "FFC7CE",
    }

    FONT_NAME = "Aptos"

    def _font(bold=False, size=10.0, colour=None):
        kw = dict(name=FONT_NAME, size=size, bold=bold)
        if colour:
            kw["color"] = colour
        return Font(**kw)

    def _fill(hex_colour):
        # Ensure full ARGB 8-char hex (openpyxl needs FFRRGGBB not RRGGBB)
        hx = hex_colour.lstrip("#")
        if len(hx) == 6:
            hx = "FF" + hx
        return PatternFill("solid", fgColor=hx)

    def _aln(horizontal="center", wrap=False):
        return Alignment(horizontal=horizontal, vertical="center",
                         wrap_text=wrap)

    def _border(left="thin", right="thin", top="thin", bottom="thin"):
        def s(style): return Side(style=style) if style else Side(style=None)
        return Border(left=s(left), right=s(right), top=s(top), bottom=s(bottom))

    # Border presets
    OG_BORDER    = _border(left="medium", right="thin",   top="medium",  bottom="medium")
    TOC_BORDER   = _border(left="medium", right="thin",   top="thin",    bottom="thin")
    COLHDR_INNER = _border(left="thin",   right="thin",   top="thin",    bottom="thin")
    COLHDR_FIRST = _border(left="medium", right="thin",   top="thin",    bottom="thin")
    COLHDR_LAST  = _border(left="thin",   right="medium", top="thin",    bottom="thin")
    DATA_INNER   = _border(left="thin",   right="thin",   top="thin",    bottom="thin")
    DATA_FIRST   = _border(left="medium", right="thin",   top="thin",    bottom="thin")
    DATA_LAST    = _border(left="thin",   right="medium", top="thin",    bottom="thin")

    def _set(cell, value=None, font=None, fill=None, alignment=None, border=None):
        if value is not None:
            cell.value = value
        if font:      cell.font      = font
        if fill:      cell.fill      = fill
        if alignment: cell.alignment = alignment
        if border:    cell.border    = border

    def _write_row(ws, row_num, values, font, fill, borders, heights):
        """Write 6-cell row with per-cell borders."""
        ws.row_dimensions[row_num].height = heights
        for ci, (val, bdr) in enumerate(zip(values, borders), start=1):
            _set(ws.cell(row=row_num, column=ci),
                 value=val, font=font, fill=fill,
                 alignment=_aln("center", wrap=(heights == 18)),
                 border=bdr)

    # ═══════════════════════════════════════════════════════════
    # SHEET 1 — TIS Accreditation Status
    # ═══════════════════════════════════════════════════════════
    ws1 = wb.active
    ws1.title = "TIS Accreditation Status"
    ws1.freeze_panes = "A1"

    # Column widths (from reference)
    ws1.column_dimensions["A"].width = 46.0
    ws1.column_dimensions["B"].width = 14.0
    ws1.column_dimensions["D"].width = 22.0
    ws1.column_dimensions["E"].width = 10.0
    ws1.column_dimensions["F"].width = 16.0

    COL_HDRS = ["System", "Machine Type ID", "Version", "State", "Count", "CoA Expiry"]

    row_num = 1

    for app_letter, og_display in APPENDIX_MAP:
        og_data = peak_df[peak_df["Owning Group"] == og_display]
        if og_data.empty:
            continue

        # ── Owning Group header (merged A:F) ──────────────────────
        ws1.merge_cells(
            start_row=row_num, start_column=1,
            end_row=row_num,   end_column=6
        )
        og_label = og_display
        _set(ws1.cell(row=row_num, column=1),
             value=og_label,
             font=_font(bold=True, size=11.0, colour=OG_FONT),
             fill=_fill(OG_FILL),
             alignment=_aln("center", wrap=True),
             border=OG_BORDER)
        # Apply borders to cells B-F of merged row too
        for ci in range(2, 7):
            c = ws1.cell(row=row_num, column=ci)
            c.fill   = _fill(OG_FILL)
            c.border = OG_BORDER
        ws1.row_dimensions[row_num].height = 18.0
        row_num += 1

        for toc in sorted(og_data["TOC / LPC"].dropna().unique()):
            toc_data = og_data[og_data["TOC / LPC"] == toc]

            # ── TOC sub-header (merged A:F) ──────────────────────
            ws1.merge_cells(
                start_row=row_num, start_column=1,
                end_row=row_num,   end_column=6
            )
            _set(ws1.cell(row=row_num, column=1),
                 value=toc,
                 font=_font(bold=True, size=10.0, colour=TOC_FONT),
                 fill=_fill(TOC_FILL),
                 alignment=_aln("center", wrap=True),
                 border=TOC_BORDER)
            for ci in range(2, 7):
                c = ws1.cell(row=row_num, column=ci)
                c.fill   = _fill(TOC_FILL)
                c.border = TOC_BORDER
            ws1.row_dimensions[row_num].height = 15.0
            row_num += 1

            # ── Column headers ────────────────────────────────────
            col_hdr_borders = [COLHDR_FIRST] + [COLHDR_INNER] * 4 + [COLHDR_LAST]
            ws1.row_dimensions[row_num].height = 15.0
            for ci, (hdr, bdr) in enumerate(zip(COL_HDRS, col_hdr_borders), start=1):
                _set(ws1.cell(row=row_num, column=ci),
                     value=hdr,
                     font=_font(bold=True, size=10.0),
                     fill=_fill(COL_FILL),
                     alignment=_aln("center", wrap=True),
                     border=bdr)
            row_num += 1

            # ── Data rows ─────────────────────────────────────────
            for _, drow in toc_data.sort_values(
                    ["System", "Machine Type ID", "Version"]).iterrows():
                state    = str(drow.get("State", ""))
                fill_hex = STATE_FILL.get(state, "FFFFFF")

                coa = drow.get("CoA Expiry", "")

                vals    = [
                    drow.get("System",          ""),
                    drow.get("Machine Type ID", ""),
                    drow.get("Version",         ""),
                    state,
                    drow.get("Count", 0),
                    coa,
                ]
                data_borders = [DATA_FIRST] + [DATA_INNER] * 4 + [DATA_LAST]
                ws1.row_dimensions[row_num].height = 14.0
                for ci, (val, bdr) in enumerate(zip(vals, data_borders), start=1):
                    _set(ws1.cell(row=row_num, column=ci),
                         value=val,
                         font=_font(bold=False, size=10.0),
                         fill=_fill(fill_hex),
                         alignment=_aln("center"),
                         border=bdr)
                row_num += 1

        row_num += 1  # spacer blank row between OG groups

    # ═══════════════════════════════════════════════════════════
    # SHEET 2 — Pivot Data
    # ═══════════════════════════════════════════════════════════
    ws2 = wb.create_sheet("Pivot Data")

    PIVOT_COLS   = ["Owning Group", "TOC / LPC", "System",
                    "Machine Type ID", "Version", "State", "Count", "CoA Expiry"]
    PIVOT_WIDTHS = {"A": 38.0, "B": 34.0, "C": 46.0, "D": 15.0,
                    "E": 14.0, "F": 24.0, "G": 10.0, "H": 16.0}
    # Left-aligned columns (from reference)
    LEFT_COLS_PIVOT = {1, 2, 3}  # Owning Group, TOC/LPC, System

    for cl, w in PIVOT_WIDTHS.items():
        ws2.column_dimensions[cl].width = w

    # Header row
    for ci, hdr in enumerate(PIVOT_COLS, start=1):
        _set(ws2.cell(row=1, column=ci),
             value=hdr,
             font=_font(bold=True, size=10.0),
             fill=_fill(COL_FILL),
             alignment=_aln("center"),
             border=_border())

    # Data rows
    export_df = peak_df[[c for c in PIVOT_COLS if c in peak_df.columns]]
    for ri, (_, row) in enumerate(export_df.iterrows(), start=2):
        state    = str(row.get("State", ""))
        fill_hex = STATE_FILL.get(state, "FFFFFF")
        for ci, col in enumerate(PIVOT_COLS, start=1):
            val  = row.get(col, "")
            halign = "left" if ci in LEFT_COLS_PIVOT else "center"
            _set(ws2.cell(row=ri, column=ci),
                 value=val,
                 font=_font(bold=False, size=10.0),
                 fill=_fill(fill_hex),
                 alignment=Alignment(horizontal=halign, vertical="center"),
                 border=_border())

    ws2.freeze_panes = "A2"

    # ═══════════════════════════════════════════════════════════
    # SHEET 3 — Accreditation Status Chart
    # ═══════════════════════════════════════════════════════════
    ws3 = wb.create_sheet("Accreditation Status Chart")
    ws3.column_dimensions["A"].width = 28.0
    ws3.column_dimensions["B"].width = 12.0

    STATE_ORDER = ["Accredited", "Pilot phase", "Accreditation expired",
                   "Application acknowledged"]
    counts = {
        s: int(peak_df[peak_df["State"] == s]["Count"].sum())
        for s in STATE_ORDER
        if int(peak_df[peak_df["State"] == s]["Count"].sum()) > 0
    }
    grand_total = sum(counts.values())

    # Header
    for ci, hdr in enumerate(["Accreditation Status", "Count"], start=1):
        _set(ws3.cell(row=1, column=ci),
             value=hdr,
             font=_font(bold=True, size=10.0),
             fill=_fill(COL_FILL),
             alignment=_aln("center"),
             border=_border())

    # State rows
    for ri, (state, cnt) in enumerate(counts.items(), start=2):
        fill_hex = STATE_FILL.get(state, "FFFFFF")
        for ci, val in enumerate([state, cnt], start=1):
            _set(ws3.cell(row=ri, column=ci),
                 value=val,
                 font=_font(bold=False, size=10.0),
                 fill=_fill(fill_hex),
                 alignment=_aln("center"),
                 border=_border())

    # Total row
    last = 1 + len(counts)
    for ci, val in enumerate(["Total Devices", grand_total], start=1):
        _set(ws3.cell(row=last+1, column=ci),
             value=val,
             font=_font(bold=True, size=10.0))

    # ═══════════════════════════════════════════════════════════
    # NATIVE PIE CHART (openpyxl) — linked to Sheet 3 data cells
    # Matches reference exactly: PieChart, varyColors, per-slice colours,
    # category + % labels, legend right, title "Accreditation Status"
    # ═══════════════════════════════════════════════════════════
    from openpyxl.chart import PieChart, Reference
    from openpyxl.chart.series import DataPoint
    from openpyxl.chart.label import DataLabelList
    from openpyxl.drawing.colors import ColorChoice
    from openpyxl.chart.shapes import GraphicalProperties
    from openpyxl.drawing.line import LineProperties
    from openpyxl.chart.legend import Legend

    # Number of state rows (excludes header + Total Devices row)
    n_states = len(counts)
    last_state_row = 1 + n_states   # rows 2..n_states+1

    pie = PieChart()
    pie.title = "Accreditation Status"
    pie.varyColors = True
    pie.firstSliceAng = 0

    # Data: counts column (B2:Bn)
    data_ref = Reference(ws3, min_col=2, min_row=1, max_row=last_state_row)
    pie.add_data(data_ref, titles_from_data=True)

    # Categories: state names column (A2:An)
    cat_ref = Reference(ws3, min_col=1, min_row=2, max_row=last_state_row)
    pie.set_categories(cat_ref)

    # Per-slice fill colours — match state colours exactly
    SLICE_COLOURS = {
        "Accredited":               "C6EFCE",
        "Pilot phase":              "FFEB9C",
        "Accreditation expired":    "FFC7CE",
        "Application acknowledged": "FFEB9C",
    }
    state_list = list(counts.keys())
    for idx, state in enumerate(state_list):
        hex_col = SLICE_COLOURS.get(state, "CCCCCC")
        dp = DataPoint(idx=idx)
        dp.graphicalProperties = GraphicalProperties(
            solidFill=hex_col
        )
        pie.series[0].dPt.append(dp)

    # Data labels: show category name + percentage
    pie.dLbls = DataLabelList()
    pie.dLbls.showCatName  = True
    pie.dLbls.showPercent  = True
    pie.dLbls.showSerName  = True
    pie.dLbls.showVal      = False
    pie.dLbls.showLegendKey = True
    pie.dLbls.showLeaderLines = True

    # Legend on the right, overlaid
    pie.legend = Legend()
    pie.legend.legendPos = "r"
    pie.legend.overlay   = True

    # Size (EMUs): 6480000 x 5040000 = 6.75" x 5.25" ≈ 15cm x 7.5cm
    pie.width  = 15
    pie.height = 7.5

    # Position: anchor at col=3 (D), row=1 (row 2, 0-indexed)
    from openpyxl.drawing.spreadsheet_drawing import OneCellAnchor, AnchorMarker
    from openpyxl.drawing.xdr import XDRPositiveSize2D

    marker = AnchorMarker(col=3, colOff=0, row=1, rowOff=0)
    size   = XDRPositiveSize2D(cx=6480000, cy=5040000)
    pie.anchor = OneCellAnchor(_from=marker, ext=size)

    ws3.add_chart(pie)
    print("    ✓ Native PieChart added to 'Accreditation Status Chart' sheet")

    # ── Save ──────────────────────────────────────────────────
    fname = f"FRACC - TIS Accreditation Status - Latest - {timestamp}.xlsx"
    fpath = os.path.join(output_dir, fname)
    wb.save(fpath)
    print(f"    Saved FRACC Excel → {fname}")
    return fpath, counts, grand_total


def generate_pie_chart(counts: dict, total: int, output_dir: str) -> str:
    """
    Render the accreditation status pie chart at 180 DPI.
    Returns the saved PNG path.
    """
    print("\n[4] Generating pie chart…")

    COLOUR_MAP = {
        "Accredited":               "#00B050",
        "Pilot phase":              "#92D050",
        "Accreditation expired":    "#EE0000",
        "Application acknowledged": "#FFC000",
    }

    labels = list(counts.keys())
    values = list(counts.values())
    colors = [COLOUR_MAP.get(l, "#AAAAAA") for l in labels]

    fig, ax = plt.subplots(figsize=(9, 5.5))
    fig.patch.set_facecolor("white")

    wedges, _ = ax.pie(
        values,
        colors=colors,
        startangle=140,
        wedgeprops=dict(linewidth=1.2, edgecolor="white"),
        radius=0.88,
    )

    # Labels outside wedges
    for wedge, val in zip(wedges, values):
        angle = (wedge.theta2 + wedge.theta1) / 2
        pct   = val / total * 100
        x  = 0.62 * np.cos(np.radians(angle))
        y  = 0.62 * np.sin(np.radians(angle))
        lx = 1.18 * np.cos(np.radians(angle))
        ly = 1.18 * np.sin(np.radians(angle))
        ax.annotate(
            f"{pct:.1f}%\n({val:,})",
            xy=(x, y), xytext=(lx, ly),
            ha="center", va="center",
            fontsize=9.5, fontweight="bold", color="#333333",
            arrowprops=dict(arrowstyle="-", color="#888888", lw=0.8),
            bbox=dict(boxstyle="round,pad=0.2", facecolor="white",
                      edgecolor="none", alpha=0.85),
        )

    # Legend — right side
    legend_patches = [
        mpatches.Patch(
            facecolor=c, edgecolor="#cccccc", linewidth=0.8,
            label=f"{l}  ({v:,})",
        )
        for l, v, c in zip(labels, values, colors)
    ]
    legend = ax.legend(
        handles=legend_patches,
        loc="center left",
        bbox_to_anchor=(1.02, 0.5),
        frameon=True, framealpha=1.0, edgecolor="#cccccc",
        fontsize=10,
        title="Accreditation Status", title_fontsize=10.5,
        labelspacing=1.0, handlelength=1.5, handleheight=1.2,
    )
    legend.get_frame().set_linewidth(0.8)

    ax.set_title(
        f"TIS Accreditation Status Summary\n(Total: {total:,} devices)",
        fontsize=13, fontweight="bold", color="#1F3864", pad=12,
    )

    plt.tight_layout(rect=[0, 0, 0.78, 1])
    pie_path = os.path.join(output_dir, "pie_chart_fracc.png")
    plt.savefig(pie_path, dpi=180, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close()
    print(f"    Saved pie chart → {os.path.basename(pie_path)}")
    return pie_path


# ════════════════════════════════════════════════════════════════════
# STEP 5 — Patch the Word document
# ════════════════════════════════════════════════════════════════════

def _set_highlight(run, colour="yellow"):
    rpr = run._r.get_or_add_rPr()
    hl  = rpr.find(qn("w:highlight"))
    if hl is None:
        hl = OxmlElement("w:highlight")
        rpr.append(hl)
    hl.set(qn("w:val"), colour)


def _add_bookmark(para, name, bm_id="199"):
    existing = [b.get(qn("w:name"))
                for b in para._p.findall(".//" + qn("w:bookmarkStart"))]
    if name in existing:
        return
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(bm_id))
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(bm_id))
    para._p.insert(0, bm_start)
    para._p.append(bm_end)


def _make_hyperlink(para, text, anchor):
    """
    Replace paragraph content with a single internal hyperlink (w:anchor).
    Preserves existing run formatting and applies the Hyperlink character style.
    """
    existing_runs = para._p.findall(qn("w:r"))
    if existing_runs:
        existing_rpr = existing_runs[0].find(qn("w:rPr"))
        rpr = copy.deepcopy(existing_rpr) if existing_rpr is not None \
              else OxmlElement("w:rPr")
    else:
        rpr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rpr.insert(0, rStyle)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    r = OxmlElement("w:r")
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    hyperlink.append(r)

    for child in list(para._p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            para._p.remove(child)
    para._p.append(hyperlink)


def patch_word_document(
    template_path:   str,
    paper_date_str:  str,
    pie_chart_path:  str,
    news_image_path: str,
    output_dir:      str,
) -> str:
    """
    Apply all patches to the Word template and save the final paper.

    Patches:
      1. Meeting Date paragraph → 'XX Month' with yellow highlight
      2. Paper Date paragraph   → paper_date_str
      3. Update para bold date  → paper_date_str
      4. Trailer table date     → 'XX Month' with yellow highlight
      5. Pie chart image (image1.png in zip) replaced
      6. News image (image2.png in zip) replaced  [if supplied]
      7. Appendix list entries (A–P) → internal Word hyperlinks
    """
    print(f"\n[5] Patching Word document…")
    doc = Document(template_path)

    # ── Parse paper date ────────────────────────────────────────────
    for fmt in ("%d %B %Y", "%d/%m/%Y", "%-d %B %Y"):
        try:
            paper_dt = datetime.strptime(paper_date_str.strip(), fmt)
            break
        except ValueError:
            continue
    else:
        paper_dt = datetime.now()
        print(f"    WARNING: Could not parse '{paper_date_str}', "
              f"using today ({paper_dt.strftime('%d %B %Y')})")

    paper_date_display = paper_dt.strftime("%-d %B %Y")   # e.g. "16 June 2026"
    paper_month_day    = paper_dt.strftime("%-d %B")       # e.g. "16 June"
    paper_year         = f" {paper_dt.year}"               # e.g. " 2026"

    # ── Locate key paragraphs ───────────────────────────────────────
    paras = doc.paragraphs
    meeting_date_para = next(
        (i for i, p in enumerate(paras) if p.text.strip().startswith("Meeting Date:")),
        None,
    )
    paper_date_para = next(
        (i for i, p in enumerate(paras) if p.text.strip().startswith("Paper Date:")),
        None,
    )
    update_para = next(
        (i for i, p in enumerate(paras)
         if "this report was run on" in p.text or "report was run on" in p.text),
        None,
    )
    # First "Appendix A…" paragraph before END OF PAPER
    toc_list_start = next(
        (i for i, p in enumerate(paras) if p.text.strip().startswith("Appendix A")),
        None,
    )

    print(f"    meeting_date_para={meeting_date_para}, "
          f"paper_date_para={paper_date_para}, "
          f"update_para={update_para}, "
          f"toc_list_start={toc_list_start}")

    MONTH_NAMES = [
        "January","February","March","April","May","June",
        "July","August","September","October","November","December",
    ]

    # ── Patch 1: Meeting Date ────────────────────────────────────────
    if meeting_date_para is not None:
        p = paras[meeting_date_para]
        for run in p.runs:
            if any(m in run.text for m in MONTH_NAMES + ["XX Month", "Month"]):
                run.text = "XX Month"
                _set_highlight(run, "yellow")
                break
        print("    ✓ Meeting Date → 'XX Month' (highlighted)")

    # ── Patch 2: Paper Date ─────────────────────────────────────────
    if paper_date_para is not None:
        p = paras[paper_date_para]
        for run in p.runs:
            if any(m in run.text for m in MONTH_NAMES + ["2025","2026","2027","2028"]):
                run.text = paper_date_display
                break
        print(f"    ✓ Paper Date → '{paper_date_display}'")

    # ── Patch 3: Update paragraph bold date ─────────────────────────
    if update_para is not None:
        p    = paras[update_para]
        runs = p.runs
        for ri, run in enumerate(runs):
            if run.bold and any(m in run.text for m in MONTH_NAMES):
                run.text = paper_month_day
                # Update the immediately following bold run (year part)
                if ri + 1 < len(runs) and runs[ri + 1].bold:
                    runs[ri + 1].text = paper_year
                break
        print(f"    ✓ Update para date → '{paper_month_day}{paper_year}'")

    # ── Patch 4: Trailer table date cell ────────────────────────────
    if doc.tables:
        trailer_tbl = doc.tables[-1]
        try:
            date_cell = trailer_tbl.rows[0].cells[1]
            date_para = date_cell.paragraphs[0]
            for run in date_para.runs:
                run.clear()
            if date_para.runs:
                r      = date_para.runs[0]
                r.text = "XX Month"
                r.bold = True
                _set_highlight(r, "yellow")
                for extra in date_para.runs[1:]:
                    extra._r.getparent().remove(extra._r)
            print("    ✓ Trailer table date → 'XX Month' (highlighted)")
        except (IndexError, AttributeError):
            print("    WARNING: Could not patch trailer table date cell")

    # ── Patch 5 & 6: Replace images in the zip ──────────────────────
    doc.save(template_path + ".tmp.docx")

    with zipfile.ZipFile(template_path + ".tmp.docx", "r") as zin:
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if (item.filename == "word/media/image1.png"
                        and pie_chart_path and os.path.exists(pie_chart_path)):
                    with open(pie_chart_path, "rb") as f:
                        zout.writestr(item, f.read())
                elif (item.filename == "word/media/image2.png"
                        and news_image_path and os.path.exists(news_image_path)):
                    with open(news_image_path, "rb") as f:
                        zout.writestr(item, f.read())
                else:
                    zout.writestr(item, zin.read(item.filename))
        buf.seek(0)

    os.remove(template_path + ".tmp.docx")

    # ── Patch 7: Appendix list → hyperlinks ─────────────────────────
    doc2  = Document(buf)
    paras2 = doc2.paragraphs

    # Ensure Appendix I bookmark exists (may be missing in older templates)
    for p in paras2:
        if "Appendix I" in p.text and "GTS Elizabeth Line" in p.text:
            _add_bookmark(p, "Appendix_I", bm_id="199")

    # Full list of appendix letters A–P (note: no J)
    ALL_LETTERS = [a for a, _ in APPENDIX_MAP] + ["M", "N", "O", "P"]
    APPENDIX_NAMES = {a: n for a, n in APPENDIX_MAP}
    APPENDIX_NAMES.update({
        "M": "Accreditation Status",
        "N": "RDG Compliance Standards - Recent Updates",
        "O": "Third Party Retailer Systems",
        "P": "Governance Lifecycle Tracking Grid",
    })

    toc_start    = None
    converted    = 0
    toc_finished = False

    for i, p in enumerate(paras2):
        if toc_finished:
            break
        txt = p.text.strip()

        # Detect start of TOC appendix list
        if toc_start is None and txt.startswith("Appendix A"):
            toc_start = i

        if toc_start is not None and i >= toc_start:
            for letter in ALL_LETTERS:
                if txt.startswith(f"Appendix {letter}") and BOOKMARK_MAP.get(letter):
                    _make_hyperlink(p, txt, BOOKMARK_MAP[letter])
                    converted += 1
                    if letter == "P":
                        toc_finished = True
                    break

    print(f"    ✓ Converted {converted} appendix entries to internal hyperlinks")

    # ── Patch 8: Strip "Appendix X — / - " prefix from section heading paragraphs ──
    import re as _re
    _strip_pat = _re.compile(
        r"^Appendix\s+[A-Z]\s*[-\u2013\u2014]\s*",
        _re.IGNORECASE
    )
    _headings_stripped = 0
    for _p in doc2.paragraphs:
        _txt = _p.text.strip()
        if _strip_pat.match(_txt):
            _new_text = _strip_pat.sub("", _p.text.lstrip()).rstrip()
            if _p.runs:
                _p.runs[0].text = _new_text
                for _r in _p.runs[1:]:
                    _r.text = ""
                _headings_stripped += 1
    print(f"    ✓ Stripped Appendix prefix from {_headings_stripped} section headings in Word doc")

    # ── Save final document ──────────────────────────────────────────
    out_name = (f"FRACC Paper - Accreditation Status Update "
                f"- {paper_date_display}.docx")
    out_path = os.path.join(output_dir, out_name)

    final_buf = io.BytesIO()
    doc2.save(final_buf)
    final_buf.seek(0)
    with open(out_path, "wb") as f:
        f.write(final_buf.read())

    print(f"    Saved Word document → {out_name}")
    return out_path


# ════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Generate the monthly FRACC TIS Accreditation Status paper.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--collated", required=True,
        help="Path to the collated TOC TIS master Excel report "
             "(raw stacked or pivot export format).",
    )
    parser.add_argument(
        "--template", required=True,
        help="Path to the FRACC Paper Word template (.docx).",
    )
    parser.add_argument(
        "--paper-date", required=True,
        help='Paper date, e.g. "16 June 2026" or "16/06/2026".',
    )
    parser.add_argument(
        "--output-dir", default=".",
        help="Output folder for generated files (default: current directory).",
    )
    parser.add_argument(
        "--news-image", default=None,
        help="(Optional) Path to ASSIST News widget screenshot (.png) to embed.",
    )

    args = parser.parse_args()

    if not os.path.exists(args.collated):
        print(f"ERROR: Collated report not found: {args.collated}")
        sys.exit(1)
    if not os.path.exists(args.template):
        print(f"ERROR: Word template not found: {args.template}")
        sys.exit(1)

    os.makedirs(args.output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")

    print("=" * 62)
    print("  FRACC TIS Accreditation Status Report Generator")
    print("=" * 62)
    print(f"  Collated report : {args.collated}")
    print(f"  Template        : {args.template}")
    print(f"  Paper date      : {args.paper_date}")
    print(f"  Output dir      : {args.output_dir}")
    print("=" * 62)

    # ── Pipeline ────────────────────────────────────────────────────
    df      = load_collated_report(args.collated)
    peak_df = apply_peak_counts(df)

    excel_path, state_counts, grand_total = build_fracc_excel(
        peak_df, args.paper_date, args.output_dir, timestamp,
    )
    pie_path  = generate_pie_chart(state_counts, grand_total, args.output_dir)
    docx_path = patch_word_document(
        template_path   = args.template,
        paper_date_str  = args.paper_date,
        pie_chart_path  = pie_path,
        news_image_path = args.news_image,
        output_dir      = args.output_dir,
    )

    print("\n" + "=" * 62)
    print("  ✓  Generation complete")
    print("=" * 62)
    print(f"  Excel → {os.path.basename(excel_path)}")
    print(f"  Word  → {os.path.basename(docx_path)}")
    print(f"  Chart → {os.path.basename(pie_path)}")
    print("=" * 62)


if __name__ == "__main__":
    main()
